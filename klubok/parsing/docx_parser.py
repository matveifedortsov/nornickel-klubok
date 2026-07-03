"""DOCX -> Document(chunks). Тот же контракт, что и pdf_parser.parse_pdf.

Корпус трека на треть состоит из .docx (внутренние статьи/обзоры) — без этого
парсера они просто не попадают в пайплайн. Тяжёлая библиотека (python-docx)
импортируется лениво, как и в pdf_parser.py.
"""
from __future__ import annotations

from pathlib import Path

from klubok.ontology import Document
from klubok.parsing.pdf_parser import chunk_text, _doc_id


def parse_docx(path: str | Path, max_chars: int = 1200) -> Document:
    """Распарсить один DOCX в Document. Требует python-docx (ленивый импорт)."""
    import docx  # python-docx

    path = Path(path)
    doc_id = _doc_id(path)
    document = Document(doc_id=doc_id, title=path.stem, source_path=str(path))

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    text = "\n\n".join(parts)
    document.chunks.extend(chunk_text(text, doc_id, page=None, max_chars=max_chars))
    return document


def parse_dir(folder: str | Path, pattern: str = "*.docx") -> list[Document]:
    folder = Path(folder)
    return [parse_docx(p) for p in sorted(folder.glob(pattern))]
